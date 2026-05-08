# -*- coding: utf-8 -*-
"""
PAP MediaRoom Generator
Wewnętrzne narzędzie redakcyjne do tworzenia informacji prasowych
na podstawie materiałów dostarczonych przez klientów.
"""

import streamlit as st
import anthropic
import json
import re
import io
from docx import Document
import pdfplumber

# ============================================================
# KONFIGURACJA STRONY
# ============================================================

st.set_page_config(
    page_title="PAP MediaRoom Generator",
    page_icon="📰",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ============================================================
# SYSTEM PROMPT (zasady PAP MediaRoom)
# ============================================================

SYSTEM_PROMPT = """Jesteś redaktorem PAP MediaRoom. Tworzysz informacje prasowe według standardów PAP MediaRoom.

ZASADY KLUCZOWE:
1. STRUKTURA: tytuł z oznaczeniem (MediaRoom) na końcu, pogrubiony lead (2 do 4 zdań, kto/co/gdzie/kiedy z atrybucją), akapit pomostowy, korpus z piramidą odwróconą, akapit kontekstowy, stopka.
2. TYTUŁ: format "Podmiot: konkretna teza" lub "Wydarzenie: konkretna teza". Bez sloganów, bez wykrzykników, bez wartościujących przymiotników.
3. JĘZYK: rzeczowy, neutralny, bez wartościowania poza cytatami. Bez metafor, sloganów, hipersuperlatywów. Bez "my/nasz" poza cytatami. UNIKAĆ STRONY BIERNEJ w miarę możliwości - preferować stronę czynną z atrybucją sprawcy (zamiast "zostało ogłoszone" pisać "ministerstwo ogłosiło"). Dopuszczalna strona bierna z atrybucją w opisach badań i ustaleń ("wykazano, że", "badanie objęło").
4. CYTATY: wprowadzane zróżnicowanymi czasownikami (powiedział, podkreślił, ocenił, zaznaczył, dodał, wyjaśnił, zauważył, wskazał). Atrybucja: imię nazwisko, funkcja. Cudzysłowy polskie „".
5. LICZBY: każda dana ma atrybucję. Format: "51 proc.", "12 mln zł", "5,7 mld zł", "19 kwietnia", "2026 r."
6. AKRONIMY: pełna nazwa + (skrót) przy pierwszym użyciu.
7. STOPKA STAŁA (zawsze kończy informację, bez modyfikacji):
"Źródło informacji: PAP MediaRoom

UWAGA: Za materiał opublikowany przez redakcję PAP MediaRoom odpowiedzialność ponosi jego nadawca, wskazany każdorazowo jako „źródło informacji"."

ZAKAZANE:
- Wprowadzanie informacji spoza materiału wejściowego (chyba że tryb uzupełnień = contextual).
- Wymyślanie cytatów lub modyfikowanie ich treści.
- Zmiana funkcji/tytułów rozmówców.
- Slogany, język marketingowy, wykrzykniki, pytania retoryczne.

ZADANIE:
Otrzymujesz materiał wejściowy od klienta podzielony na ponumerowane zdania. Twoim zadaniem jest:
1. Stworzyć informację prasową w stylu PAP MediaRoom.
2. Dla każdego zdania wejściowego sklasyfikować je: USED (wykorzystane), EXCLUDED (kategoria wykluczona z liczenia: kontakt prasowy, boilerplate, stopki, klauzule), SKIPPED (świadomie pominięte z krótkim uzasadnieniem).
3. Podzielić gotową informację prasową na zdania i dla każdego zdania wskazać, które zdania wejścia je wspierają (lista ID), oraz oznaczyć czy zdanie jest dodane (added=true) jako uzupełnienie kontekstowe spoza materiału (tylko w trybie contextual).
4. Wystawić ostrzeżenia jeśli materiał jest ubogi, sprzeczny lub nie nadaje się do publikacji.

Wynik zwróć przez wywołanie narzędzia `generate_press_release` z odpowiednimi parametrami."""


# Schema narzędzia, które wymusza strukturę odpowiedzi modelu
GENERATE_TOOL = {
    "name": "generate_press_release",
    "description": "Zwraca wygenerowaną informację prasową lub depeszę (zależnie od wskazanego stylu) wraz z mapowaniem wykorzystania materiału wejściowego i listą ostrzeżeń.",
    "input_schema": {
        "type": "object",
        "properties": {
            "press_release_text": {
                "type": "string",
                "description": "Pełny tekst informacji prasowej w stylu PAP MediaRoom, z łamaniami linii między akapitami."
            },
            "press_release_sentences": {
                "type": "array",
                "description": "Wygenerowana informacja prasowa podzielona na pojedyncze zdania.",
                "items": {
                    "type": "object",
                    "properties": {
                        "text": {
                            "type": "string",
                            "description": "Treść pojedynczego zdania informacji prasowej."
                        },
                        "added": {
                            "type": "boolean",
                            "description": "Czy zdanie jest dodaniem spoza materiału wejściowego (tylko w trybie contextual)."
                        },
                        "supported_by": {
                            "type": "array",
                            "items": {"type": "integer"},
                            "description": "Lista ID zdań wejściowych, które wspierają to zdanie. Pusta lista oznacza brak podstawy w materiale (czerwona flaga)."
                        }
                    },
                    "required": ["text", "added", "supported_by"]
                }
            },
            "input_mapping": {
                "type": "array",
                "description": "Klasyfikacja każdego zdania wejściowego.",
                "items": {
                    "type": "object",
                    "properties": {
                        "id": {
                            "type": "integer",
                            "description": "Numer zdania wejściowego."
                        },
                        "status": {
                            "type": "string",
                            "enum": ["USED", "EXCLUDED", "SKIPPED"],
                            "description": "USED gdy wykorzystane, EXCLUDED gdy z kategorii wykluczonej, SKIPPED gdy świadomie pominięte."
                        },
                        "reason": {
                            "type": "string",
                            "description": "Krótkie uzasadnienie statusu (np. 'kontakt prasowy', 'powtórzenie informacji'). Dla USED może być pusty string."
                        }
                    },
                    "required": ["id", "status", "reason"]
                }
            },
            "warnings": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Lista ostrzeżeń dla operatora (np. materiał ubogi, sprzeczność wewnętrzna, treść reklamowa). Pusta lista jeśli brak."
            }
        },
        "required": ["press_release_text", "press_release_sentences", "input_mapping", "warnings"]
    }
}


# ============================================================
# SYSTEM PROMPT I TOOL DLA TRYBU STAŻYSTY
# ============================================================

SYSTEM_PROMPT_TRAINEE = """Jesteś doświadczonym redaktorem PAP MediaRoom. Otrzymujesz informację prasową napisaną przez stażystę PAP. Twoim zadaniem jest ocenić ją i poprawić zgodnie ze standardami PAP MediaRoom, zachowując wszystkie fakty z oryginału stażysty.

ZASADY KLUCZOWE PAP MEDIAROOM:
1. STRUKTURA: tytuł z oznaczeniem (MediaRoom) na końcu, pogrubiony lead (2 do 4 zdań, kto/co/gdzie/kiedy z atrybucją), akapit pomostowy, korpus z piramidą odwróconą, akapit kontekstowy, stopka.
2. TYTUŁ: format "Podmiot: konkretna teza" lub "Wydarzenie: konkretna teza". Bez sloganów, bez wykrzykników, bez wartościujących przymiotników.
3. JĘZYK: rzeczowy, neutralny, bez wartościowania poza cytatami. Bez metafor, sloganów, hipersuperlatywów. Bez "my/nasz" poza cytatami. UNIKAĆ STRONY BIERNEJ w miarę możliwości - preferować stronę czynną z atrybucją sprawcy (zamiast "zostało ogłoszone" pisać "ministerstwo ogłosiło"). Dopuszczalna strona bierna z atrybucją w opisach badań i ustaleń ("wykazano, że", "badanie objęło").
4. CYTATY: wprowadzane zróżnicowanymi czasownikami (powiedział, podkreślił, ocenił, zaznaczył, dodał, wyjaśnił, zauważył, wskazał). Atrybucja: imię nazwisko, funkcja. Cudzysłowy polskie „".
5. LICZBY: każda dana ma atrybucję. Format: "51 proc.", "12 mln zł", "5,7 mld zł", "19 kwietnia", "2026 r."
6. AKRONIMY: pełna nazwa + (skrót) przy pierwszym użyciu.
7. STOPKA STAŁA (zawsze kończy informację, bez modyfikacji):
"Źródło informacji: PAP MediaRoom

UWAGA: Za materiał opublikowany przez redakcję PAP MediaRoom odpowiedzialność ponosi jego nadawca, wskazany każdorazowo jako „źródło informacji"."

ZAKAZANE:
- Wymyślanie faktów spoza tekstu stażysty. Pracujesz wyłącznie na tym, co napisał stażysta.
- Modyfikowanie treści cytatów (wolno tylko poprawiać atrybucję i cudzysłowy).
- Zmiana funkcji/tytułów rozmówców.
- Slogany, język marketingowy, wykrzykniki, pytania retoryczne.
- Cichych poprawek bez uzasadnienia. Każda zmiana musi być wyjaśniona w liście uchybień.
- Pomijania drobnych uchybień stylistycznych. Stażysta uczy się także na drobiazgach.

ZADANIE:
1. Stwórz poprawioną wersję tekstu zgodną ze standardami PAP MediaRoom. Zachowaj formatowanie akapitów - oddzielaj akapity pustą linią.
2. Wskaż konkretne uchybienia (co zostało źle zrobione i wymagało poprawki). Każde uchybienie sklasyfikuj według kategorii i wagi (poważne / drobne).
3. Wskaż konkretne zgodności ze stylebookiem (co stażysta zrobił dobrze, godne pochwały). To ważny element feedbacku motywacyjnego.

ZASADA SPÓJNOŚCI - KRYTYCZNA:
- Każda zmiana, którą wprowadzasz w tekście względem oryginału stażysty, MUSI mieć odpowiadające jej uchybienie w liście "issues". Bez wyjątków.
- Jeśli zmieniasz zdanie ze strony biernej na czynną, to jest uchybienie kategorii "język" (waga: drobne) - wymień je.
- Jeśli przeformułowujesz zdanie dla lepszej rzeczowości, to jest uchybienie - wymień je.
- Jeśli poprawiasz cudzysłowy, atrybucję, format liczb - to są uchybienia, wymień każde.
- Jeśli zmieniasz strukturę (przesuwasz akapity, dzielisz długie zdania) - to są uchybienia, wymień każde.
- Liczba uchybień powinna być co najmniej tak duża, jak liczba zmienionych zdań. Możesz wymienić więcej, jeśli jedno zdanie zawierało kilka problemów.
- Jedyna sytuacja, w której lista uchybień może być pusta: gdy nie wprowadzasz ŻADNYCH zmian w tekście stażysty (tekst jest perfekcyjny w obecnej formie).

KATEGORIE UCHYBIEŃ I ZGODNOŚCI:
- struktura (lead, korpus, kolejność, stopka)
- język (rzeczowość, neutralność, brak marketingu)
- cytaty (atrybucja, cudzysłowy, czasowniki wprowadzające)
- liczby (atrybucja, format)
- tytuł
- lead
- inne

Wynik zwróć przez wywołanie narzędzia review_trainee_text."""


TRAINEE_TOOL = {
    "name": "review_trainee_text",
    "description": "Ocenia i poprawia tekst (informację prasową lub depeszę) napisany przez stażystę PAP, zwracając poprawioną wersję, listę uchybień oraz listę elementów zgodnych ze stylebookiem.",
    "input_schema": {
        "type": "object",
        "properties": {
            "corrected_text": {
                "type": "string",
                "description": "Pełny poprawiony tekst informacji prasowej, gotowy do publikacji. Zachowaj formatowanie akapitów (puste linie między akapitami). Zachowaj wszystkie fakty z oryginału stażysty - nie wymyślaj nowych."
            },
            "compliances": {
                "type": "array",
                "description": "Lista konkretnych elementów, które stażysta zrobił zgodnie ze stylebookiem PAP MediaRoom (feedback pozytywny).",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {
                            "type": "string",
                            "enum": ["struktura", "język", "cytaty", "liczby", "tytuł", "lead", "inne"],
                            "description": "Kategoria zgodności."
                        },
                        "description": {
                            "type": "string",
                            "description": "Konkretny opis tego, co zostało zrobione dobrze."
                        }
                    },
                    "required": ["category", "description"]
                }
            },
            "issues": {
                "type": "array",
                "description": "Lista uchybień znalezionych w tekście stażysty (feedback wymagający poprawki).",
                "items": {
                    "type": "object",
                    "properties": {
                        "category": {
                            "type": "string",
                            "enum": ["struktura", "język", "cytaty", "liczby", "tytuł", "lead", "inne"],
                            "description": "Kategoria uchybienia."
                        },
                        "severity": {
                            "type": "string",
                            "enum": ["poważne", "drobne"],
                            "description": "Waga uchybienia: poważne (np. brak stopki, brak leadu, błędna atrybucja) lub drobne (np. niepolskie cudzysłowy, mała literówka stylistyczna)."
                        },
                        "description": {
                            "type": "string",
                            "description": "Konkretny opis uchybienia oraz tego, co zostało poprawione."
                        },
                        "fragment": {
                            "type": "string",
                            "description": "Fragment z oryginału stażysty, którego dotyczy uchybienie. Pusty string jeśli uchybienie dotyczy całości."
                        }
                    },
                    "required": ["category", "severity", "description", "fragment"]
                }
            }
        },
        "required": ["corrected_text", "compliances", "issues"]
    }
}


# ============================================================
# SYSTEM PROMPTY I TOOL DLA STYLU DEPESZA PAP
# ============================================================

DEPESZA_STYLE_RULES = """ZASADY DEPESZY PAP:

1. ZWIĘZŁOŚĆ I PRECYZJA: tekst możliwie najkrótszy. Z każdego zdania powinno dać się coś ująć i tekst nadal będzie zrozumiały. Bez metafor, porównań, dygresji. Konkrety i fakty, uporządkowane od najważniejszych.

2. TYTUŁ: streszcza depeszę, nie jest zagadką. Bez pytań, bez niedopowiedzeń. UWAGA: depesza NIE zawiera oznaczenia (MediaRoom) - to jest treść redakcyjna, nie materiał komercyjny.

3. LID: pierwsze 1 do 2 zdań zawiera wszystkie kluczowe informacje: kto, co, kiedy, gdzie, dlaczego. Jeśli czytelnik przeczyta tylko lid, powinien znać sens wydarzenia.

4. PIRAMIDA ODWRÓCONA: najpierw najważniejsze fakty, potem dopowiedzenia i tło. Nigdy odwrotnie. Pozwala to redakcjom skracać depeszę „od dołu" bez utraty sensu.

5. JĘZYK NEUTRALNY, BEZ EMOCJI: nie oceniamy, nie komentujemy, nie podsumowujemy. Bez przymiotników oceniających („słuszny", „szokujący", „tragiczny"). Bez sugerowania emocji („niestety", „na szczęście"). Bez spekulacji. Mówimy, co się stało, nie co o tym myślimy.

6. CYTATY I ŹRÓDŁA: cytaty dosłowne, z podaniem źródła. Bez parafrazowania zmieniającego sens. Cytaty krótkie i istotne. Jeśli osoba wypowiadająca się jest istotna dla sprawy, podajemy jej funkcję.

7. OSZCZĘDNOŚĆ JĘZYKOWA: bez powtórzeń (nie piszemy dwa razy tego samego innymi słowami). Bez oczywistości („W poniedziałek, który był pierwszym dniem tygodnia"). Bez skomplikowanej składni. Krótkie zdania, jednoznaczne komunikaty.

8. TON I FORMA: depesza nie mówi do czytelnika. Bez zwrotów bezpośrednich, bez pytań retorycznych, bez stylizacji. Bez pierwszej osoby. Trzecia osoba, oficjalny rejestr językowy.

9. PRECYZJA I ODPOWIEDZIALNOŚĆ: każda informacja sprawdzona, każda liczba pewna, każda interpretacja ostrożna. Jeśli coś niepotwierdzone, pisać wprost: „według relacji świadków", „jak poinformowano nieoficjalnie".

10. RÓWNOWAGA STRON: wszędzie tam, gdzie to możliwe, przedstawiać stanowiska obu stron sporu, konfliktu, wydarzenia. Jeśli jedna strona stawia zarzuty, druga powinna mieć szansę na odpowiedź. Jeśli takiej odpowiedzi nie ma, zaznaczyć, np. „do czasu publikacji nie otrzymaliśmy komentarza".

11. STRONA BIERNA: unikać w miarę możliwości - preferować stronę czynną z atrybucją sprawcy (zamiast „zostało ogłoszone" pisać „ministerstwo ogłosiło"). Dopuszczalna strona bierna w opisach badań i ustaleń.

DEPESZA NIE ZAWIERA:
- Oznaczenia (MediaRoom) w tytule
- Stopki "Źródło informacji: PAP MediaRoom..."
- Tonu marketingowego, promocyjnego
- Komentarzy, ocen, spekulacji
- Zwrotów do czytelnika"""


DEPESZA_STYLE_ONLY_INSTRUCTION = """TRYB PRACY: ZMIANA TYLKO STYLU

Zachowaj kolejność informacji z materiału wejściowego. Twoim zadaniem są WYŁĄCZNIE korekty stylistyczne wymagane przez depeszowy stylebook (rzeczowość, neutralność, oszczędność, trzecia osoba, krótkie zdania, brak ozdobników). Nie oceniaj niezależnie wagi informacji ani nie zmieniaj ich kolejności w stosunku do materiału wejściowego."""


DEPESZA_PRIORITY_INSTRUCTION = """TRYB PRACY: OCENA WAGI I PIRAMIDA ODWRÓCONA

Działasz jako doświadczony dziennikarz agencyjny PAP. Twoim profesjonalnym obowiązkiem jest niezależnie ocenić wagę każdej informacji i zbudować depeszę według właściwej piramidy odwróconej. Kolejność z materiału wejściowego NIE jest wiążąca - autor mógł ułożyć informacje nieprawidłowo z perspektywy depeszy.

JAK OCENIAĆ WAGĘ INFORMACJI:
- Społeczne znaczenie (czy wydarzenie dotyczy ogółu, dużych grup, regionu, kraju)
- Liczba osób dotkniętych skutkami
- Nowość, bezprecedensowość, świeżość
- Pilność i aktualność (czy redakcje czekają na tę informację)
- Konsekwencje krótko- i długoterminowe
- Zainteresowanie czytelnika ogólnopolskiego

JAK BUDOWAĆ DEPESZĘ:
1. Przeanalizuj wszystkie fakty z materiału i zdecyduj niezależnie, co jest najważniejsze.
2. Tytuł: streszcza najważniejszą informację (nie pierwszą z materiału - najważniejszą).
3. Lid (1 do 2 zdań): kluczowe fakty (kto, co, kiedy, gdzie, dlaczego) wokół najważniejszej informacji.
4. Korpus: kolejne fakty uporządkowane od najważniejszych do najmniej ważnych.
5. Końcówka: tło, kontekst, dopowiedzenia, mniej istotne szczegóły.
6. Informacje wyraźnie poboczne (kontakty, historia firmy nieistotna dla wydarzenia, drobne dygresje, treść marketingowa) umieść na końcu albo pomiń.

ODNOTOWANIE ZMIANY KOLEJNOŚCI:
- W trybie generowania z materiału klienta: w warnings krótko wyjaśnij, jeśli znacząco zmieniłeś kolejność informacji (np. "Najważniejszą informacją było X z 4. akapitu materiału - przeniesione do leadu").
- W trybie poprawiania tekstu stażysty: w issues dodaj uchybienie kategorii "struktura" za każde znaczące przesunięcie (np. "Lead nie zawierał najważniejszej informacji - przeniesiono ją z 3. akapitu")."""


SYSTEM_PROMPT_DEPESZA = f"""Jesteś dziennikarzem agencyjnym PAP. Tworzysz depesze prasowe na podstawie materiału klienta. Depesza to podstawowa forma dziennikarskiego przekazu - krótka, rzeczowa, szybka. Jej celem nie jest opowiadanie historii, lecz błyskawiczne dostarczenie faktów redakcjom w całej Polsce.

{DEPESZA_STYLE_RULES}

ZAKAZANE:
- Wprowadzanie informacji spoza materiału wejściowego (chyba że tryb uzupełnień = contextual).
- Wymyślanie cytatów lub modyfikowanie ich treści.
- Zmiana funkcji/tytułów rozmówców.
- Slogany, marketing, wykrzykniki, pytania retoryczne, zwroty do czytelnika.

ZADANIE:
Otrzymujesz materiał wejściowy od klienta podzielony na ponumerowane zdania. Twoim zadaniem jest:
1. Stworzyć depeszę PAP w opisanym wyżej stylu (BEZ oznaczenia (MediaRoom), BEZ stopki MediaRoom).
2. Dla każdego zdania wejściowego sklasyfikować je: USED (wykorzystane), EXCLUDED (kategoria wykluczona: kontakt prasowy, boilerplate, stopki, klauzule, czysta treść marketingowa), SKIPPED (świadomie pominięte z uzasadnieniem).
3. Podzielić depeszę na zdania, wskazać ID wspierających zdań wejścia, oznaczyć added=true tylko w trybie contextual.
4. Wystawić ostrzeżenia jeśli materiał jest ubogi, wewnętrznie sprzeczny, jednostronny (np. zarzuty bez odpowiedzi drugiej strony) lub w przeważającej mierze marketingowy.

Wynik zwróć przez wywołanie narzędzia generate_press_release."""


SYSTEM_PROMPT_TRAINEE_DEPESZA = f"""Jesteś doświadczonym dziennikarzem agencyjnym PAP. Otrzymujesz depeszę napisaną przez stażystę PAP. Twoim zadaniem jest ocenić ją i poprawić zgodnie z zasadami depeszy PAP, zachowując wszystkie fakty z oryginału stażysty.

{DEPESZA_STYLE_RULES}

ZAKAZANE:
- Wymyślanie faktów spoza tekstu stażysty. Pracujesz wyłącznie na tym, co napisał stażysta.
- Modyfikowanie treści cytatów (wolno tylko poprawiać atrybucję i cudzysłowy).
- Zmiana funkcji/tytułów rozmówców.
- Cichych poprawek bez uzasadnienia. Każda zmiana musi być wyjaśniona w liście uchybień.
- Pomijania drobnych uchybień stylistycznych. Stażysta uczy się także na drobiazgach.
- Slogany, marketing, wykrzykniki, pytania retoryczne, zwroty do czytelnika.

ZADANIE:
1. Stwórz poprawioną wersję depeszy zgodną ze standardami PAP. Zachowaj formatowanie akapitów - oddzielaj akapity pustą linią. Pamiętaj, że depesza NIE zawiera oznaczenia (MediaRoom) ani stopki MediaRoom.
2. Wskaż konkretne uchybienia (co zostało źle zrobione i wymagało poprawki). Każde uchybienie sklasyfikuj według kategorii i wagi (poważne / drobne).
3. Wskaż konkretne zgodności ze stylebookiem (co stażysta zrobił dobrze, godne pochwały). To ważny element feedbacku motywacyjnego.

ZASADA SPÓJNOŚCI - KRYTYCZNA:
- Każda zmiana, którą wprowadzasz w tekście względem oryginału stażysty, MUSI mieć odpowiadające jej uchybienie w liście "issues". Bez wyjątków.
- Jeśli zmieniasz zdanie ze strony biernej na czynną, to jest uchybienie kategorii "język" (waga: drobne) - wymień je.
- Jeśli przeformułowujesz zdanie dla lepszej rzeczowości, to jest uchybienie - wymień je.
- Jeśli skracasz dygresje, eliminujesz ozdobniki, usuwasz oceny - to są uchybienia, wymień każde.
- Jeśli poprawiasz cudzysłowy, atrybucję, format liczb - to są uchybienia, wymień każde.
- Jeśli zmieniasz strukturę (porządkujesz piramidę odwróconą, dzielisz długie zdania) - to są uchybienia, wymień każde.
- Liczba uchybień powinna być co najmniej tak duża, jak liczba zmienionych zdań.
- Jedyna sytuacja, w której lista uchybień może być pusta: gdy nie wprowadzasz ŻADNYCH zmian (tekst jest perfekcyjny w obecnej formie).

KATEGORIE UCHYBIEŃ I ZGODNOŚCI:
- struktura (lid, korpus, piramida odwrócona)
- język (rzeczowość, neutralność, brak emocji, brak ozdobników)
- cytaty (atrybucja, cudzysłowy, funkcja osoby)
- liczby (precyzja, format)
- tytuł
- lid
- równowaga (przedstawienie obu stron)
- inne

Wynik zwróć przez wywołanie narzędzia review_trainee_text."""


# ============================================================
# AUTORYZACJA
# ============================================================

def check_authentication():
    """Sprawdza czy użytkownik jest zalogowany. Zwraca True jeśli tak."""
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if st.session_state.authenticated:
        return True

    # Ekran logowania
    st.title("PAP MediaRoom Generator")
    st.caption("Wewnętrzne narzędzie redakcyjne. Logowanie wymagane.")

    col_left, col_center, col_right = st.columns([1, 2, 1])
    with col_center:
        with st.form("login_form", clear_on_submit=False):
            password = st.text_input("Hasło", type="password", placeholder="wpisz hasło")
            submitted = st.form_submit_button("Zaloguj", type="primary", use_container_width=True)
            if submitted:
                expected = st.secrets.get("app_password", "")
                if password == expected and expected != "":
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("Niepoprawne hasło.")
    return False


# ============================================================
# EKSTRAKCJA TEKSTU Z PLIKÓW
# ============================================================

def extract_text_from_docx(file_bytes):
    """Wyciąga tekst z pliku Word (.docx)."""
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)
    # Wyciągnij też tekst z tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def extract_text_from_pdf(file_bytes):
    """Wyciąga tekst z pliku PDF."""
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


# ============================================================
# DZIELENIE NA ZDANIA
# ============================================================

POLISH_ABBREVIATIONS = [
    "r", "mln", "mld", "proc", "pkt", "tj", "np", "tzw", "tys",
    "str", "nr", "ul", "al", "woj", "godz", "min", "sek",
    "wg", "por", "art", "ust", "par", "dr", "prof", "inż", "mgr",
    "płk", "gen", "płd", "pn", "wsch", "zach", "ok"
]


DOT_PLACEHOLDER = "\u0001"  # tymczasowy znacznik chronionych kropek wewnątrz skrótów


def split_sentences(text):
    """Dzieli tekst na zdania, uwzględniając polskie skróty.

    Strategia: kropki wewnątrz skrótów (np. "r.", "proc.", "m.in.") oraz
    między cyframi (np. "12.04.2026") tymczasowo zastępujemy znacznikiem,
    dzielimy tekst, potem przywracamy oryginalne kropki.
    """
    cleaned = re.sub(r"\s+", " ", text.strip())
    if not cleaned:
        return []

    # Zastąp kropki w skrótach znacznikiem (case insensitive: "Dr." i "dr.")
    abbrev_pattern = r"\b(" + "|".join(POLISH_ABBREVIATIONS) + r")\."
    protected = re.sub(
        abbrev_pattern,
        lambda m: m.group(1) + DOT_PLACEHOLDER,
        cleaned,
        flags=re.IGNORECASE
    )
    # Specjalny przypadek "m.in."
    protected = re.sub(
        r"\bm\.in\.",
        "m" + DOT_PLACEHOLDER + "in" + DOT_PLACEHOLDER,
        protected,
        flags=re.IGNORECASE
    )
    # Liczby z kropką (daty, numery, dziesiętne)
    protected = re.sub(
        r"(\d)\.(\d)",
        lambda m: m.group(1) + DOT_PLACEHOLDER + m.group(2),
        protected
    )

    parts = re.findall(r"[^.!?]+[.!?]+", protected)
    if not parts:
        parts = [protected]
    return [p.replace(DOT_PLACEHOLDER, ".").strip() for p in parts if p.strip()]


# ============================================================
# WYWOŁANIE CLAUDE
# ============================================================

def build_user_prompt(sentences, format_mode, supplement_mode, excluded):
    format_label = (
        "sztywny (tytuł, lead, korpus, kontakt, stopka)"
        if format_mode == "rigid"
        else "elastyczny (zachowane zasady języka i atrybucji, dopuszczona swoboda struktury)"
    )
    supplement_label = (
        "ZEROWA TOLERANCJA - wyłącznie materiał wejściowy, żadnych dodatków spoza"
        if supplement_mode == "zero"
        else "KONTEKSTOWE - można dodać podstawowy kontekst (np. czym jest dana instytucja), wszystkie dodatki oznaczyć added=true"
    )
    excl_label = ", ".join(excluded) if excluded else "brak"

    numbered = "\n".join(f"[{i+1}] {s}" for i, s in enumerate(sentences))

    return f"""KONFIGURACJA:
- Format: {format_label}
- Tryb uzupełnień: {supplement_label}
- Kategorie wykluczone z liczenia wykorzystania (sklasyfikuj odpowiednie zdania jako EXCLUDED): {excl_label}

MATERIAŁ WEJŚCIOWY (ponumerowane zdania):
{numbered}

Wykonaj zadanie i zwróć wyłącznie JSON według podanego formatu."""


_SENTINEL = object()


def _safe_get(obj, key, default=None):
    """Pobiera wartość spod klucza/atrybutu, niezależnie od typu obiektu.
    Działa na słownikach, modelach pydantic, namedtuples, dataclassach.
    """
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(key, default)
    val = getattr(obj, key, _SENTINEL)
    if val is not _SENTINEL:
        return val
    try:
        return obj[key]
    except (TypeError, KeyError, IndexError):
        return default


def _to_plain_python(obj):
    """Konwertuje cokolwiek do czystych typów Pythona przez serializację JSON.
    Niezależnie od tego, czy SDK zwraca model pydantic, dict, MappingProxy,
    namedtuple czy inny obiekt, wynikiem jest gwarantowany dict/list/str/int/bool/None.
    """
    def _fallback(o):
        # Najpierw próbujemy model_dump (pydantic v2)
        dump = getattr(o, "model_dump", None)
        if callable(dump):
            try:
                return dump(mode="python")
            except Exception:
                pass
        # Potem .dict() (pydantic v1, jeśli ktoś tego używa)
        dict_method = getattr(o, "dict", None)
        if callable(dict_method) and not isinstance(o, dict):
            try:
                result = dict_method()
                if isinstance(result, dict):
                    return result
            except Exception:
                pass
        # Potem __dict__ (zwykłe obiekty Pythona)
        if hasattr(o, "__dict__"):
            d = vars(o)
            if d:
                return {k: v for k, v in d.items() if not k.startswith("_")}
        # Ostateczność: konwersja do stringa
        return str(o)

    # Krok 1: pre-konwersja, jeśli sam obiekt jest pydantic
    dump = getattr(obj, "model_dump", None)
    if callable(dump):
        try:
            obj = dump(mode="python")
        except Exception:
            pass

    # Krok 2: round-trip przez JSON wymusza czyste typy
    try:
        return json.loads(json.dumps(obj, default=_fallback, ensure_ascii=False))
    except Exception:
        # Awaryjny manualny przepływ rekursywny
        if isinstance(obj, dict):
            return {str(k): _to_plain_python(v) for k, v in obj.items()}
        if isinstance(obj, (list, tuple)):
            return [_to_plain_python(item) for item in obj]
        return obj


def call_claude(sentences, format_mode, supplement_mode, excluded, style="mediaroom", depesza_mode="style_only"):
    """Wywołuje Claude API używając tool use i zwraca strukturyzowany wynik.

    style: "mediaroom" (informacja prasowa MediaRoom) lub "depesza" (depesza PAP)
    depesza_mode: "style_only" (zachowaj kolejność, zmień tylko styl) lub
                  "evaluate" (oceń wagę informacji, zastosuj piramidę odwróconą)
    """
    client = anthropic.Anthropic(api_key=st.secrets["anthropic_api_key"])
    user_prompt = build_user_prompt(sentences, format_mode, supplement_mode, excluded)

    if style == "depesza":
        if depesza_mode == "evaluate":
            system_prompt = SYSTEM_PROMPT_DEPESZA + "\n\n" + DEPESZA_PRIORITY_INSTRUCTION
        else:
            system_prompt = SYSTEM_PROMPT_DEPESZA + "\n\n" + DEPESZA_STYLE_ONLY_INSTRUCTION
    else:
        system_prompt = SYSTEM_PROMPT

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=16000,
        temperature=0,
        system=system_prompt,
        tools=[GENERATE_TOOL],
        tool_choice={"type": "tool", "name": "generate_press_release"},
        messages=[{"role": "user", "content": user_prompt}]
    )

    # Sprawdź czy odpowiedź nie została obcięta limitem tokenów
    stop_reason = getattr(message, "stop_reason", None)
    if stop_reason == "max_tokens":
        st.warning(
            "Odpowiedź modelu została obcięta limitem tokenów. "
            "Wynik może być niekompletny. Spróbuj z krótszym materiałem."
        )

    # Wynik znajduje się w bloku tool_use. Normalizujemy do czystego dict.
    for block in message.content:
        if block.type == "tool_use" and block.name == "generate_press_release":
            return _to_plain_python(block.input)

    raise ValueError(
        "Model nie wywołał oczekiwanego narzędzia. "
        f"Otrzymano typy bloków: {[b.type for b in message.content]}"
    )


def call_claude_trainee(text, style="mediaroom", depesza_mode="style_only"):
    """Wywołuje Claude API w trybie poprawiania tekstu stażysty.

    style: "mediaroom" (informacja prasowa MediaRoom) lub "depesza" (depesza PAP)
    depesza_mode: "style_only" (zachowaj kolejność, zmień tylko styl) lub
                  "evaluate" (oceń wagę informacji, zastosuj piramidę odwróconą)
    """
    client = anthropic.Anthropic(api_key=st.secrets["anthropic_api_key"])

    user_prompt = f"""TEKST STAŻYSTY DO OCENY I POPRAWY:

{text}

Wykonaj zadanie i wywołaj narzędzie review_trainee_text."""

    if style == "depesza":
        if depesza_mode == "evaluate":
            system_prompt = SYSTEM_PROMPT_TRAINEE_DEPESZA + "\n\n" + DEPESZA_PRIORITY_INSTRUCTION
        else:
            system_prompt = SYSTEM_PROMPT_TRAINEE_DEPESZA + "\n\n" + DEPESZA_STYLE_ONLY_INSTRUCTION
    else:
        system_prompt = SYSTEM_PROMPT_TRAINEE

    message = client.messages.create(
        model="claude-opus-4-6",
        max_tokens=16000,
        temperature=0,
        system=system_prompt,
        tools=[TRAINEE_TOOL],
        tool_choice={"type": "tool", "name": "review_trainee_text"},
        messages=[{"role": "user", "content": user_prompt}]
    )

    # Sprawdź czy odpowiedź nie została obcięta limitem tokenów
    stop_reason = getattr(message, "stop_reason", None)
    if stop_reason == "max_tokens":
        st.warning(
            "Odpowiedź modelu została obcięta limitem tokenów. "
            "Wynik może być niekompletny. Spróbuj z krótszym tekstem."
        )

    for block in message.content:
        if block.type == "tool_use" and block.name == "review_trainee_text":
            return _to_plain_python(block.input)

    raise ValueError(
        "Model nie wywołał oczekiwanego narzędzia. "
        f"Otrzymano typy bloków: {[b.type for b in message.content]}"
    )


# ============================================================
# OBLICZENIA I RENDEROWANIE
# ============================================================

def compute_usage(mapping):
    total = len(mapping)
    excluded = sum(1 for m in mapping if _safe_get(m, "status") == "EXCLUDED")
    used = sum(1 for m in mapping if _safe_get(m, "status") == "USED")
    denominator = total - excluded
    if denominator == 0:
        return {"percent": 0, "used": 0, "total": 0, "skipped": 0}
    skipped = denominator - used
    return {
        "percent": round((used / denominator) * 100),
        "used": used,
        "total": denominator,
        "skipped": skipped
    }


def escape_html(s):
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def render_input_with_highlights(sentences, mapping):
    """Renderuje materiał wejściowy z oznaczonymi statusami zdań."""
    map_dict = {_safe_get(m, "id"): m for m in mapping if _safe_get(m, "id") is not None}
    parts = []
    for i, s in enumerate(sentences):
        sid = i + 1
        m = map_dict.get(sid, {"status": "USED", "reason": ""})
        status = _safe_get(m, "status", "USED")
        reason = _safe_get(m, "reason", "") or ""

        if status == "USED":
            style = "background: transparent;"
            tooltip = ""
        elif status == "EXCLUDED":
            style = "background: #f1efe8; color: #888; text-decoration: line-through;"
            tooltip = f' title="EXCLUDED: {escape_html(reason)}"' if reason else ' title="EXCLUDED"'
        else:  # SKIPPED
            style = "background: rgba(186, 117, 23, 0.15); color: #633806;"
            tooltip = f' title="SKIPPED: {escape_html(reason)}"' if reason else ' title="SKIPPED"'

        parts.append(
            f'<span style="{style} padding: 1px 3px; border-radius: 3px;"{tooltip}>'
            f'[{sid}] {escape_html(s)}</span>'
        )
    return " ".join(parts)


def render_output_with_flags(pr_sentences):
    """Renderuje informację prasową z oznaczonymi flagami."""
    parts = []
    for s in pr_sentences:
        text = _safe_get(s, "text", "")
        added = _safe_get(s, "added", False)
        supported_by = _safe_get(s, "supported_by", []) or []

        if added:
            style = "background: rgba(186, 117, 23, 0.18); border-bottom: 1px dashed #BA7517;"
            tooltip = ' title="Uzupełnienie spoza materiału - zweryfikuj"'
        elif not supported_by:
            style = "background: rgba(226, 75, 74, 0.18); border-bottom: 1px solid #A32D2D;"
            tooltip = ' title="Brak podstawy w materiale wejściowym - zweryfikuj"'
        else:
            style = ""
            tooltip = f' title="Oparte na zdaniach: {", ".join(str(x) for x in supported_by)}"'

        parts.append(
            f'<span style="{style} padding: 1px 3px; border-radius: 3px;"{tooltip}>'
            f'{escape_html(text)}</span>'
        )
    return " ".join(parts)


def display_results(sentences, result):
    """Wyświetla wynik generowania."""
    # Walidacja typów - bezpiecznik na wypadek gdyby model zwrócił coś niezgodnego ze schematem
    raw_input_mapping = _safe_get(result, "input_mapping", [])
    input_mapping = raw_input_mapping if isinstance(raw_input_mapping, list) else []
    input_mapping = [m for m in input_mapping if isinstance(m, dict)]

    raw_pr_sentences = _safe_get(result, "press_release_sentences", [])
    pr_sentences = raw_pr_sentences if isinstance(raw_pr_sentences, list) else []
    pr_sentences = [s for s in pr_sentences if isinstance(s, dict)]

    raw_warnings = _safe_get(result, "warnings", [])
    warnings = raw_warnings if isinstance(raw_warnings, list) else []
    warnings = [w for w in warnings if isinstance(w, str)]

    raw_pr_text = _safe_get(result, "press_release_text", "")
    pr_text = raw_pr_text if isinstance(raw_pr_text, str) else ""

    # Wykryj sytuację niezgodną ze schematem
    schema_issues = []
    if not isinstance(raw_pr_text, str):
        schema_issues.append(f"press_release_text: oczekiwano stringa, otrzymano {type(raw_pr_text).__name__}")
    if not isinstance(raw_pr_sentences, list):
        schema_issues.append(f"press_release_sentences: oczekiwano listy, otrzymano {type(raw_pr_sentences).__name__}")
    if not isinstance(raw_input_mapping, list):
        schema_issues.append(f"input_mapping: oczekiwano listy, otrzymano {type(raw_input_mapping).__name__}")

    if schema_issues:
        st.error(
            "Model zwrócił odpowiedź niezgodną ze schematem narzędzia. "
            "Wynik może być niekompletny.\n\n"
            "Szczegóły:\n" + "\n".join(f"- {s}" for s in schema_issues)
        )
        with st.expander("Diagnostyka (surowa odpowiedź modelu)"):
            st.json(result if isinstance(result, dict) else {"raw": str(result)[:5000]})

    usage = compute_usage(input_mapping)
    flagged = sum(
        1 for s in pr_sentences
        if not _safe_get(s, "added") and not (_safe_get(s, "supported_by") or [])
    )
    added = sum(1 for s in pr_sentences if _safe_get(s, "added"))

    # Statystyki
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric(
            "Wykorzystanie materiału",
            f"{usage['percent']}%",
            help=f"{usage['used']} z {usage['total']} zdań kwalifikowanych (po wykluczeniu kategorii)"
        )
    with col2:
        st.metric(
            "Dodane fragmenty",
            added,
            help="Liczba zdań dodanych poza materiałem wejściowym (tryb kontekstowy)"
        )
    with col3:
        st.metric(
            "Czerwone flagi",
            flagged,
            help="Liczba zdań wyjścia bez wyraźnej podstawy w materiale - do weryfikacji"
        )

    # Ostrzeżenia
    if warnings:
        st.warning("**Ostrzeżenia operatora:**\n\n" + "\n".join(f"- {w}" for w in warnings))

    # Dwie kolumny: materiał i informacja
    left, right = st.columns(2)
    with left:
        st.subheader("Materiał wejściowy")
        st.caption("Najedź kursorem na fragment, aby zobaczyć status")
        html_input = render_input_with_highlights(sentences, input_mapping)
        st.markdown(
            f'<div style="font-size: 14px; line-height: 1.75;">{html_input}</div>',
            unsafe_allow_html=True
        )
        st.caption("🔘 Wykluczone (przekreślone)   🔶 Pominięte (bursztynowe)")

    with right:
        st.subheader("Informacja prasowa")
        st.caption("Najedź kursorem na fragment, aby zobaczyć źródło")
        html_output = render_output_with_flags(pr_sentences)
        st.markdown(
            f'<div style="font-size: 14px; line-height: 1.75;">{html_output}</div>',
            unsafe_allow_html=True
        )
        st.caption("🔶 Dodane spoza materiału   🔴 Bez podstawy w materiale")

    st.divider()

    # Surowy tekst do skopiowania
    st.subheader("Tekst gotowy do skopiowania")
    st.code(pr_text, language=None)

    # Pobranie pliku
    st.download_button(
        label="Pobierz jako plik tekstowy",
        data=pr_text.encode("utf-8"),
        file_name="informacja_prasowa.txt",
        mime="text/plain"
    )


def _normalize_for_compare(text):
    """Normalizuje zdanie do porównania: małe litery, pojedyncze odstępy, bez końcowych znaków interpunkcyjnych."""
    return re.sub(r"\s+", " ", text.strip().lower()).rstrip(".!?,;:")


def _render_text_with_paragraphs(text):
    """Renderuje zwykły tekst zachowując akapity (oddzielone pustymi liniami)."""
    paragraphs = re.split(r"\n\s*\n", text.strip())
    parts = []
    for para in paragraphs:
        if not para.strip():
            continue
        cleaned = re.sub(r"\s+", " ", para.strip())
        parts.append(
            f"<p style='margin: 0 0 12px 0; font-size: 14px; line-height: 1.75;'>"
            f"{escape_html(cleaned)}</p>"
        )
    return "".join(parts)


def _render_corrected_with_preserved_highlights(corrected_text, original_normalized_set):
    """Renderuje poprawiony tekst z zielonym podświetleniem zdań zachowanych z oryginału."""
    paragraphs = re.split(r"\n\s*\n", corrected_text.strip())
    parts = []
    for para in paragraphs:
        if not para.strip():
            continue
        sentences = split_sentences(para)
        if not sentences:
            cleaned = re.sub(r"\s+", " ", para.strip())
            parts.append(
                f"<p style='margin: 0 0 12px 0; font-size: 14px; line-height: 1.75;'>"
                f"{escape_html(cleaned)}</p>"
            )
            continue
        spans = []
        for s in sentences:
            is_preserved = _normalize_for_compare(s) in original_normalized_set
            if is_preserved:
                style = (
                    "background: rgba(151, 196, 89, 0.35); "
                    "padding: 1px 3px; border-radius: 3px;"
                )
                tooltip = ' title="Zdanie zachowane z oryginału stażysty"'
            else:
                style = "padding: 1px 3px;"
                tooltip = ' title="Zdanie zmienione przez redakcję"'
            spans.append(
                f'<span style="{style}"{tooltip}>{escape_html(s)}</span>'
            )
        parts.append(
            f"<p style='margin: 0 0 12px 0; font-size: 14px; line-height: 1.75;'>"
            f"{' '.join(spans)}</p>"
        )
    return "".join(parts)


def display_trainee_results(original_text, result):
    """Wyświetla wynik oceny i poprawy tekstu stażysty."""
    # Walidacja typów
    raw_corrected_text = _safe_get(result, "corrected_text", "")
    corrected_text = raw_corrected_text if isinstance(raw_corrected_text, str) else ""

    raw_compliances = _safe_get(result, "compliances", [])
    compliances = raw_compliances if isinstance(raw_compliances, list) else []
    compliances = [c for c in compliances if isinstance(c, dict)]

    raw_issues = _safe_get(result, "issues", [])
    issues = raw_issues if isinstance(raw_issues, list) else []
    issues = [i for i in issues if isinstance(i, dict)]

    # Jeśli model nie zwrócił podstawowych pól, pokaż błąd zamiast pustego widoku
    if not corrected_text:
        st.error(
            "Model nie zwrócił poprawionego tekstu. Spróbuj jeszcze raz "
            "(czasem przy długich tekstach pomaga drugie podejście)."
        )
        with st.expander("Diagnostyka (surowa odpowiedź modelu)"):
            st.json(result if isinstance(result, dict) else {"raw": str(result)[:5000]})
        return

    # Porównanie zdań: znajdź te w wersji poprawionej, które są również w oryginale
    original_sentences = split_sentences(original_text)
    original_normalized = {_normalize_for_compare(s) for s in original_sentences}

    corrected_sentences = split_sentences(corrected_text)
    preserved_count = sum(
        1 for s in corrected_sentences
        if _normalize_for_compare(s) in original_normalized
    )
    changed_count = len(corrected_sentences) - preserved_count

    # Dwie kolumny: oryginał vs wersja poprawiona
    left, right = st.columns(2)
    with left:
        st.subheader("Tekst stażysty (oryginał)")
        st.caption("Wersja przed redakcją")
        st.markdown(
            _render_text_with_paragraphs(original_text),
            unsafe_allow_html=True
        )

    with right:
        st.subheader("Wersja poprawiona")
        st.caption("Na zielono: zdania zachowane z oryginału stażysty")
        st.markdown(
            _render_corrected_with_preserved_highlights(corrected_text, original_normalized),
            unsafe_allow_html=True
        )

    st.divider()

    # Liczniki pod tekstami
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric(
            "Liczba zmian",
            changed_count,
            help="Liczba zdań w wersji poprawionej, które różnią się od oryginału stażysty."
        )
    with col2:
        st.metric(
            "Liczba uchybień",
            len(issues),
            help="Liczba elementów wymagających poprawki według stylebooka PAP MediaRoom."
        )
    with col3:
        st.metric(
            "Zrobione dobrze",
            len(compliances),
            help="Liczba elementów już zgodnych ze stylebookiem PAP MediaRoom."
        )

    st.divider()

    # Ostrzeżenie o niezgodności logicznej: jeśli zmian jest więcej niż uchybień,
    # to model nie wymienił wszystkich powodów swoich poprawek
    if changed_count > 0 and len(issues) < changed_count:
        st.warning(
            f"Model wprowadził {changed_count} zmian, ale wymienił tylko {len(issues)} uchybień. "
            f"Niektóre poprawki nie mają uzasadnienia w liście. Spróbuj wygenerować ponownie - "
            "model czasem opuszcza drobne uchybienia stylistyczne."
        )

    # Listy zgodności i uchybień
    col_good, col_bad = st.columns(2)
    with col_good:
        st.subheader("Zrobione dobrze")
        if compliances:
            for c in compliances:
                cat = _safe_get(c, "category", "inne")
                desc = _safe_get(c, "description", "")
                if cat or desc:
                    st.markdown(f"✓ **{cat.capitalize()}**: {desc}")
        else:
            st.caption("Brak wyróżnionych zgodności w tekście.")

    with col_bad:
        st.subheader("Uchybienia do poprawy")
        if issues:
            for i in issues:
                cat = _safe_get(i, "category", "inne")
                sev = _safe_get(i, "severity", "drobne")
                desc = _safe_get(i, "description", "")
                fragment = _safe_get(i, "fragment", "")
                badge = "🔴" if sev == "poważne" else "🟡"
                st.markdown(f"{badge} **{cat.capitalize()}** ({sev}): {desc}")
                if fragment:
                    st.caption(f'Fragment: „{fragment}"')
        else:
            st.caption("Brak uchybień. Tekst zgodny ze stylebookiem.")

    st.divider()

    # Surowy tekst do skopiowania
    st.subheader("Tekst gotowy do skopiowania")
    st.code(corrected_text, language=None)

    # Pobranie pliku
    st.download_button(
        label="Pobierz poprawiony tekst",
        data=corrected_text.encode("utf-8"),
        file_name="informacja_prasowa_poprawiona.txt",
        mime="text/plain",
        key="download_trainee"
    )


# ============================================================
# GŁÓWNY INTERFEJS
# ============================================================

def main_app():
    # Pasek górny
    col_title, col_logout = st.columns([5, 1])
    with col_title:
        st.title("PAP MediaRoom Generator")
    with col_logout:
        st.write("")
        if st.button("Wyloguj", use_container_width=True):
            st.session_state.authenticated = False
            st.rerun()

    # Wybór trybu pracy
    mode = st.radio(
        "Tryb pracy",
        ["Pisanie z materiału klienta", "Poprawianie tekstu stażysty"],
        horizontal=True,
        key="work_mode",
        help="Pisanie: tworzysz informację prasową na podstawie materiału od klienta. Poprawianie: oceniasz i redagujesz gotową informację prasową napisaną przez stażystę."
    )

    st.divider()

    if mode == "Pisanie z materiału klienta":
        _client_material_flow()
    else:
        _trainee_correction_flow()


def _client_material_flow():
    """Tryb pisania informacji prasowej na podstawie materiału klienta."""
    # Sidebar z ustawieniami
    with st.sidebar:
        st.header("Ustawienia generowania")
        st.caption("Tryb: pisanie z materiału klienta")

        style_label = st.radio(
            "Styl",
            ["Informacja prasowa MediaRoom", "Depesza PAP (ZROB DEPESZE)"],
            help="MediaRoom: informacja prasowa z oznaczeniem (MediaRoom) i stopką PAP MediaRoom. Depesza: krótka, neutralna depesza dziennikarska bez oznaczenia komercyjnego.",
            key="client_style"
        )
        style = "depesza" if style_label.startswith("Depesza") else "mediaroom"

        depesza_mode = "style_only"
        if style == "depesza":
            depesza_mode_label = st.radio(
                "Tryb depeszy",
                ["Zmień tylko styl", "Sprawdź wagę informacji i zmień"],
                help="Tylko styl: zachowuje kolejność z materiału, zmienia jedynie język na depeszowy. Ocena wagi: AI niezależnie ocenia wagę każdej informacji i stosuje piramidę odwróconą - może przesunąć kluczową informację z dalszej części materiału do leadu.",
                key="client_depesza_mode"
            )
            depesza_mode = "evaluate" if depesza_mode_label.startswith("Sprawdź") else "style_only"

        st.divider()

        format_label = st.radio(
            "Format informacji",
            ["Sztywny", "Elastyczny"],
            help="Sztywny: pełna struktura zgodna ze stylebookiem. Elastyczny: dopuszczalne odstępstwa od struktury.",
            key="client_format"
        )
        format_mode = "rigid" if format_label == "Sztywny" else "flexible"

        st.divider()

        supplement_label = st.radio(
            "Tryb uzupełnień",
            ["Zerowa tolerancja", "Kontekstowe"],
            help="Zerowa: tylko materiał klienta. Kontekstowe: model może dodać podstawowy kontekst, oznaczony w wyjściu.",
            key="client_supplement"
        )
        supplement_mode = "zero" if supplement_label == "Zerowa tolerancja" else "contextual"

        st.divider()

        st.subheader("Wykluczone z liczenia")
        st.caption("Te kategorie nie są liczone do współczynnika wykorzystania")
        excluded = []
        if st.checkbox("Kontakt prasowy", value=True, key="client_excl_kontakt"):
            excluded.append("kontakt prasowy")
        if st.checkbox("Nota o spółce / boilerplate", value=True, key="client_excl_nota"):
            excluded.append("nota o spółce / boilerplate")
        if st.checkbox("Klauzule prawne / disclaimer", value=True, key="client_excl_klauzule"):
            excluded.append("klauzule prawne / disclaimer")
        if st.checkbox("Stopki i podpisy", value=True, key="client_excl_stopki"):
            excluded.append("stopki i podpisy")

    # Główna część
    st.subheader("Materiał wejściowy od klienta")

    tab_text, tab_file = st.tabs(["Wklej tekst", "Wgraj plik (Word lub PDF)"])

    material_text = ""

    with tab_text:
        material_text = st.text_area(
            "Wklej tutaj materiał od klienta",
            height=300,
            label_visibility="collapsed",
            placeholder="Wklej tutaj materiał od klienta...",
            key="client_material_input"
        )

    with tab_file:
        uploaded = st.file_uploader(
            "Wybierz plik Word (.docx) lub PDF",
            type=["docx", "pdf"],
            label_visibility="collapsed",
            key="client_file_upload"
        )
        if uploaded is not None:
            try:
                file_bytes = uploaded.read()
                if uploaded.name.lower().endswith(".docx"):
                    extracted = extract_text_from_docx(file_bytes)
                else:
                    extracted = extract_text_from_pdf(file_bytes)

                if not extracted.strip():
                    st.warning("Nie udało się wyciągnąć tekstu z pliku. Sprawdź czy plik nie jest skanem.")
                else:
                    st.success(f"Wyciągnięto {len(extracted)} znaków z pliku {uploaded.name}")
                    material_text = st.text_area(
                        "Wyodrębniony tekst (możesz go edytować przed wygenerowaniem)",
                        value=extracted,
                        height=300,
                        key="client_extracted_text"
                    )
            except Exception as e:
                st.error(f"Błąd odczytu pliku: {e}")

    # Przycisk generowania
    button_label = "Generuj depeszę" if style == "depesza" else "Generuj informację prasową"
    spinner_label = "Generuję depeszę" if style == "depesza" else "Generuję informację prasową"

    if st.button(button_label, type="primary", key="client_generate"):
        if not material_text or not material_text.strip():
            st.error("Wklej najpierw materiał albo wgraj plik.")
        else:
            sentences = split_sentences(material_text)
            if len(sentences) < 2:
                st.error("Materiał zbyt krótki. Wymagane minimum 2 zdania.")
            else:
                with st.spinner(f"{spinner_label} ({len(sentences)} zdań do analizy)..."):
                    try:
                        result = call_claude(
                            sentences, format_mode, supplement_mode, excluded,
                            style=style, depesza_mode=depesza_mode
                        )
                        st.session_state.last_client_result = result
                        st.session_state.last_client_sentences = sentences
                        st.session_state.last_client_style = style
                        st.session_state.last_client_depesza_mode = depesza_mode
                    except Exception as e:
                        st.error(f"Błąd: {e}")
                        return

    # Wyświetl ostatni wynik (jeśli jest)
    if st.session_state.get("last_client_result") and st.session_state.get("last_client_sentences"):
        st.divider()
        display_results(
            st.session_state.last_client_sentences,
            st.session_state.last_client_result
        )


def _trainee_correction_flow():
    """Tryb poprawiania informacji prasowej napisanej przez stażystę."""
    # Sidebar z ustawieniami
    with st.sidebar:
        st.header("Ustawienia oceny")
        st.caption("Tryb: poprawianie tekstu stażysty")

        style_label = st.radio(
            "Styl",
            ["Informacja prasowa MediaRoom", "Depesza PAP (ZROB DEPESZE)"],
            help="MediaRoom: oceniaj wg stylebooka PAP MediaRoom (z oznaczeniem (MediaRoom) i stopką). Depesza: oceniaj wg stylebooka depeszy PAP (krótka, neutralna, bez oznaczeń komercyjnych).",
            key="trainee_style"
        )
        style = "depesza" if style_label.startswith("Depesza") else "mediaroom"

        depesza_mode = "style_only"
        if style == "depesza":
            depesza_mode_label = st.radio(
                "Tryb depeszy",
                ["Zmień tylko styl", "Sprawdź wagę informacji i zmień"],
                help="Tylko styl: zachowuje kolejność z oryginału stażysty, poprawia jedynie język. Ocena wagi: AI niezależnie ocenia wagę informacji i stosuje piramidę odwróconą - może przesunąć kluczowe informacje. Każde przesunięcie pojawi się w liście uchybień.",
                key="trainee_depesza_mode"
            )
            depesza_mode = "evaluate" if depesza_mode_label.startswith("Sprawdź") else "style_only"

        st.divider()

        st.caption(
            "Wklej lub wgraj tekst napisany przez stażystę. "
            "Aplikacja oceni go zgodnie ze wskazanym stylebookiem, "
            "wskaże uchybienia oraz zgodności i zaproponuje poprawioną wersję."
        )

    # Główna część
    subject = "depeszy" if style == "depesza" else "informacji prasowej"
    st.subheader(f"Tekst {subject} stażysty do oceny i poprawy")

    tab_text, tab_file = st.tabs(["Wklej tekst", "Wgraj plik (Word lub PDF)"])

    trainee_text = ""

    with tab_text:
        trainee_text = st.text_area(
            "Wklej tutaj tekst stażysty",
            height=300,
            label_visibility="collapsed",
            placeholder="Wklej tutaj informację prasową napisaną przez stażystę...",
            key="trainee_text_input"
        )

    with tab_file:
        uploaded = st.file_uploader(
            "Wybierz plik Word (.docx) lub PDF",
            type=["docx", "pdf"],
            label_visibility="collapsed",
            key="trainee_file_upload"
        )
        if uploaded is not None:
            try:
                file_bytes = uploaded.read()
                if uploaded.name.lower().endswith(".docx"):
                    extracted = extract_text_from_docx(file_bytes)
                else:
                    extracted = extract_text_from_pdf(file_bytes)

                if not extracted.strip():
                    st.warning("Nie udało się wyciągnąć tekstu z pliku. Sprawdź czy plik nie jest skanem.")
                else:
                    st.success(f"Wyciągnięto {len(extracted)} znaków z pliku {uploaded.name}")
                    trainee_text = st.text_area(
                        "Wyodrębniony tekst (możesz go edytować przed sprawdzeniem)",
                        value=extracted,
                        height=300,
                        key="trainee_extracted_text"
                    )
            except Exception as e:
                st.error(f"Błąd odczytu pliku: {e}")

    # Przycisk
    if st.button("Sprawdź i popraw", type="primary", key="trainee_review"):
        if not trainee_text or not trainee_text.strip():
            st.error("Wklej najpierw tekst stażysty albo wgraj plik.")
        elif len(trainee_text.strip()) < 100:
            st.error("Tekst zbyt krótki. Wymagane minimum 100 znaków.")
        else:
            with st.spinner("Sprawdzam i poprawiam tekst..."):
                try:
                    result = call_claude_trainee(trainee_text, style=style, depesza_mode=depesza_mode)
                    st.session_state.last_trainee_result = result
                    st.session_state.last_trainee_text = trainee_text
                    st.session_state.last_trainee_style = style
                    st.session_state.last_trainee_depesza_mode = depesza_mode
                except Exception as e:
                    st.error(f"Błąd: {e}")
                    return

    # Wyświetl ostatni wynik
    if st.session_state.get("last_trainee_result") and st.session_state.get("last_trainee_text"):
        st.divider()
        display_trainee_results(
            st.session_state.last_trainee_text,
            st.session_state.last_trainee_result
        )


# ============================================================
# URUCHOMIENIE
# ============================================================

if check_authentication():
    main_app()
